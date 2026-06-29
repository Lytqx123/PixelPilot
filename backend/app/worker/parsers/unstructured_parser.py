# 常规文档解析：支持PDF/DOCX/XLSX等格式，Unstructured提取文本

import logging
import os
from typing import List, Dict, Any, Optional

from .base import BaseParser

logger = logging.getLogger(__name__)


class UnstructuredParser(BaseParser):
    """
    使用 LangChain Unstructured 库进行文档解析

    支持：PDF, DOCX, XLSX, MD, TXT, CSV
    """

    # 支持的编码列表（按优先级）
    ENCODINGS = ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]

    def __init__(self):
        self._initialized = False
        self._loader_cls = None

    def _ensure_init(self):
        """延迟初始化，导入可能较重的依赖"""
        if self._initialized:
            return

        try:
            from langchain_community.document_loaders import UnstructuredFileLoader
            self._loader_cls = UnstructuredFileLoader
            self._initialized = True
            logger.info("Unstructured 解析器初始化成功")
        except ImportError as e:
            logger.warning(f"Unstructured 库未安装: {e}")
            self._initialized = False

    def parse(self, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """
        解析常规文档：无损文本提取和语义分块
        """
        logger.info(f"[UnstructuredParser] 解析 {file_path} (type={file_type})")

        # 1. 尝试 Unstructured
        text_content = self._extract_with_unstructured(file_path, file_type)

        # 2. docx 专属降级：使用 python-docx 直接提取文本
        if not text_content and file_type == "docx":
            text_content = self._extract_docx_text(file_path)

        # 3. pdf 专属降级：使用 PyPDF2 直接提取文本
        if not text_content and file_type == "pdf":
            text_content = self._extract_pdf_text(file_path)

        # 4. xlsx 专属降级：使用 openpyxl 直接提取文本
        if not text_content and file_type == "xlsx":
            text_content = self._extract_xlsx_text(file_path)

        # 5. 降级方案：直接读取文本
        if not text_content:
            text_content = self._fallback_text_extract(file_path, file_type)

        # 6. 最终降级：元数据
        if not text_content:
            return self._fallback_metadata(file_path, file_type)

        # 7. 验证提取的内容
        if len(text_content.strip()) < 10:
            logger.warning(f"提取内容过少，尝试其他方案: {file_path}")
            alt_content = self._extract_as_plain_text(file_path)
            if alt_content and len(alt_content.strip()) > len(text_content.strip()):
                text_content = alt_content

        # 8. 分块
        chunks = self.chunk_text(text_content, chunk_size=1000, overlap=200)

        if not chunks:
            return self._fallback_metadata(file_path, file_type)

        return [
            {
                "content": chunk,
                "page_number": 1,
                "paragraph": f"{file_type.upper()} 文档内容",
            }
            for chunk in chunks
        ]

    def _extract_with_unstructured(self, file_path: str, file_type: str) -> Optional[str]:
        """使用 Unstructured 提取文本"""
        self._ensure_init()
        if not self._initialized or self._loader_cls is None:
            return None

        try:
            # 根据文件类型选择加载模式
            mode = "elements" if file_type in ("pdf", "docx", "doc") else "single"
            loader = self._loader_cls(file_path, mode=mode)
            docs = loader.load()

            if not docs:
                return None

            # 合并所有元素文本
            text_parts = []
            for doc in docs:
                if hasattr(doc, "page_content") and doc.page_content:
                    content = doc.page_content.strip()
                    if content:
                        text_parts.append(content)

            return "\n".join(text_parts) if text_parts else None

        except Exception as e:
            logger.warning(f"Unstructured 提取失败: {e}")
            return None

    def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """使用 python-docx 提取 docx 文本（纯 Python 实现，无需系统依赖）"""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx 未安装，无法提取 docx 文本")
            return None

        try:
            doc = Document(file_path)
            parts = []

            # 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 保留标题层级信息
                    if para.style and para.style.name and para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading", "").strip()
                        prefix = "#" * min(int(level), 6) if level.isdigit() else "##"
                        parts.append(f"{prefix} {text}")
                    else:
                        parts.append(text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            result = "\n".join(parts)
            if result.strip():
                logger.info(f"python-docx 成功提取 {len(parts)} 个文本段落: {file_path}")
                return result
            else:
                logger.warning(f"python-docx 提取到空内容: {file_path}")
                return None

        except Exception as e:
            logger.warning(f"python-docx 提取失败: {e}")
            return None

    def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """使用 PyPDF2 提取 PDF 文本（纯 Python 实现，无需系统依赖）"""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.warning("PyPDF2 未安装，无法提取 PDF 文本")
            return None

        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)

            if total_pages == 0:
                logger.warning(f"PDF 文件无页面: {file_path}")
                return None

            parts = []
            # 提取文档信息
            if reader.metadata:
                title = reader.metadata.get("/Title", "")
                if title:
                    parts.append(f"# {title}")

            # 逐页提取文本
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    parts.append(f"## 第 {i + 1} 页\n{text.strip()}")

            result = "\n\n".join(parts)
            if result.strip():
                logger.info(f"PyPDF2 成功提取 {total_pages} 页 PDF 文本: {file_path}")
                return result
            else:
                logger.warning(f"PyPDF2 提取到空内容（可能是扫描版 PDF）: {file_path}")
                return None

        except Exception as e:
            logger.warning(f"PyPDF2 提取失败: {e}")
            return None

    def _extract_xlsx_text(self, file_path: str) -> Optional[str]:
        """使用 openpyxl 提取 xlsx 文本（纯 Python 实现，无需系统依赖）"""
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.warning("openpyxl 未安装，无法提取 xlsx 文本")
            return None

        try:
            wb = load_workbook(file_path, data_only=True)
            parts = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"## 工作表: {sheet_name}")

                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue

                # 过滤完全空行
                valid_rows = [r for r in rows if any(c is not None for c in r)]
                if not valid_rows:
                    continue

                # 表头行
                headers = [str(h) if h is not None else "" for h in valid_rows[0]]
                parts.append(" | ".join(headers))
                parts.append(" | ".join(["---"] * len(headers)))

                # 数据行
                for row in valid_rows[1:]:
                    parts.append(" | ".join([str(c) if c is not None else "" for c in row]))

                parts.append("")

            result = "\n".join(parts)
            if result.strip():
                logger.info(f"openpyxl 成功提取 {len(wb.sheetnames)} 个工作表: {file_path}")
                return result
            else:
                logger.warning(f"openpyxl 提取到空内容: {file_path}")
                return None

        except Exception as e:
            logger.warning(f"openpyxl 提取失败: {e}")
            return None

    def _fallback_text_extract(self, file_path: str, file_type: str) -> Optional[str]:
        """降级方案：直接读取文本文件"""
        if file_type not in ("txt", "md", "csv", "log"):
            return None

        # 尝试多种编码
        for encoding in self.ENCODINGS:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                if content and content.strip():
                    logger.info(f"成功使用 {encoding} 编码读取: {file_path}")
                    return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"读取文件失败 ({encoding}): {e}")
                break

        return None

    def _extract_as_plain_text(self, file_path: str) -> Optional[str]:
        """通用文本提取"""
        # 尝试作为二进制读取然后解码
        try:
            with open(file_path, "rb") as f:
                raw = f.read()

            for encoding in self.ENCODINGS:
                try:
                    text = raw.decode(encoding)
                    if text.strip():
                        return text
                except (UnicodeDecodeError, UnicodeError):
                    continue

            # 忽略解码错误
            text = raw.decode("utf-8", errors="ignore")
            return text if text.strip() else None

        except Exception as e:
            logger.warning(f"通用文本提取失败: {e}")
            return None

    def _fallback_metadata(self, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """最终降级：返回文件元数据（仅在所有解析方式都失败时调用）"""
        try:
            import os
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            content = (
                f"[{file_type.upper()}] {os.path.basename(file_path)}\n"
                f"文件大小: {file_size / 1024:.1f} KB\n"
            )
        except Exception:
            content = f"[{file_type.upper()}] {file_path}"

        return [{
            "content": content,
            "page_number": 1,
            "paragraph": f"{file_type.upper()} 文档",
        }]
